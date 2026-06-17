# utils/doc_writer.py
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_BREAK
import logging
import markdown
from bs4 import BeautifulSoup, NavigableString
from io import BytesIO

from editor.models.posts import UnifiedPostTimelineService
from editor.models.stories import StoryService
from editor.models.storyChars import StoryCharsService, StoryWithCharactersService
from editor.models.chapters import ChapterService, ChapterCharService
from editor.models.chars import CharService
from editor.utils.formatCharacter import buildChar

from flask_login import current_user

def process_inline(node, paragraph):
    #    Process inline formatting within a paragraph.

    if isinstance(node, NavigableString):
        paragraph.add_run(str(node))
        return

    if node.name == "strong":
        run = paragraph.add_run(node.get_text())
        run.bold = True

    elif node.name == "em":
        run = paragraph.add_run(node.get_text())
        run.italic = True

    else:
        for child in node.children:
            process_inline(child, paragraph)


def process_block(node, doc):
    
    # Process block-level elements.

    if node.name == "p":
        para = doc.add_paragraph()

        for child in node.children:
            process_inline(child, para)

    elif node.name == "h1":
        doc.add_heading(node.get_text(), level=1)

    elif node.name == "h2":
        doc.add_heading(node.get_text(), level=2)

    elif node.name == "h3":
        doc.add_heading(node.get_text(), level=3)

    elif node.name == "ol":

        for li in node.find_all("li", recursive=False):

            para = doc.add_paragraph(style="List Number")

            for child in li.children:

                if getattr(child, "name", None) == "p":

                    # Copy paragraph contents into list item
                    for grandchild in child.children:
                        process_inline(grandchild, para)

                else:
                    process_inline(child, para)

    elif node.name == "ul":

        for li in node.find_all("li", recursive=False):

            para = doc.add_paragraph(style="List Bullet")

            for child in li.children:

                if getattr(child, "name", None) == "p":

                    for grandchild in child.children:
                        process_inline(grandchild, para)

                else:
                    process_inline(child, para)

def markdown_to_docx_paragraph(doc: Document, input_text: str):
    # Convert markdown to HTML then parse
    html=markdown.markdown(input_text)
    soup = BeautifulSoup(html, 'html.parser')

    for node in soup.contents:

        if getattr(node, "name", None):
            process_block(node, doc)

    return

def insert_image(doc: Document, image_b64: str, mime_type: str) -> None:
    try:
        if mime_type not in ["image/png", "image/jpeg", "image/jpg"]:
            raise ValueError(f"Unsupported MIME type: {mime_type}")

        #img_data = base64.b64decode(image_b64)
        img_io = BytesIO(image_b64)
        doc.add_picture(img_io, height=Inches(4.5))  # or your preferred width

    except Exception as e:
        doc.add_paragraph(f"[Image failed to render: {e}]")

def generate_doc_from_posts(story_id, chapter_id=None) -> BytesIO:
    if chapter_id:
        posts=UnifiedPostTimelineService.get_all_posts_raw(story_id, chapter_id)
    else:
        posts=UnifiedPostTimelineService.get_all_posts_raw(story_id)
    logging.debug("Entering print generation")
    doc = Document()
    story = StoryService.get_user_story(story_id, current_user.id)

    if chapter_id:
        posts=UnifiedPostTimelineService.get_all_posts_raw(story_id, chapter_id)
        chapter = ChapterService.get_chapter(chapter_id)
        doc.add_heading(f"Book: {story.title}", level=0)
        print_chapter_intro(doc, story, chapter)
        print_posts(doc, posts)
        print_chapter_summary(doc, chapter)

    elif story.book:
        chapters=ChapterService.get_chapters(story_id)
        doc.add_heading(f"Title {story.title}", level=0)
        doc.add_heading(f"Author {current_user.user_name}", level=2)
        print_story_chars(doc, story)
        for chapter in chapters:
            posts=UnifiedPostTimelineService.get_all_posts_raw(story_id, chapter.chapter_id)
            print_chapter_intro(doc, story, chapter)
            print_posts(doc, posts)
            print_chapter_summary(doc, chapter)
    else:
        doc.add_heading(f"Title {story.title}", level=0)
        doc.add_heading(f"Author {current_user.user_name}", level=2)
        print_story_chars(doc, story)

        posts=UnifiedPostTimelineService.get_all_posts_raw(story_id)
        print_posts(doc, posts)

    if story.note != "":
        doc.add_heading("Story Notes:", level=2)
        markdown_to_docx_paragraph(doc, story.note)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def print_posts(doc, posts):
    # Print posts will work for both stories and chapters
    ix=1
    if not posts:
        doc.add_paragraph("No content yet.")
        return
    
    char_print=False
    for post in posts:
        logging.debug(f"Printing post {ix}")
        if post.source=="post" and post.creator != "user":
            title = f"Post {ix}"
            ix +=1 
            if char_print:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                char_print=False
            doc.add_heading(title, level=2)
        
        if post.creator == "model":
            markdown_to_docx_paragraph(doc, post.content)

        if post.part_type == "text":
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            doc.add_heading("Character", level=2)
            markdown_to_docx_paragraph(doc, post.part_text)
            char_print=True
        
        if post.part_image_mime_type != "" and post.part_image_mime_type != None :  # presence implies image exists
            insert_image(doc, post.part_image_data, post.part_image_mime_type)

    return

def print_chapter_intro(doc, story, chapter):

    doc.add_heading(f"Chapter {chapter.position}: {chapter.title}", level=1)
    doc.add_heading(f"Author {current_user.user_name}", level=2)
    
    doc.add_heading("Introduction", level=2)
    markdown_to_docx_paragraph(doc, chapter.introduction)
    
    firstTime=True
    chapter_chars=ChapterCharService.get_chapter_chars(story.story_id, chapter.chapter_id)
    for chapter_char in chapter_chars:
        if firstTime:
            doc.add_heading("Characters", level=2)
            firstTime=False
        char=CharService.get_character(chapter_char.char_id, current_user.id)
        if chapter_char.override:
            note=chapter_char.note
        else:
            base_note=StoryCharsService.get_story_chars_base(story.story_id,chapter_char.char_id)
            note=base_note.note

        if story.book:
            char_title=f"**Character:** {char.name}"
            markdown_to_docx_paragraph(doc, char_title)
            if chapter_char.override:
                note=chapter_char.note
            else:
                base_note=StoryCharsService.get_story_chars_base(story.story_id,chapter_char.char_id)
                note=base_note.note
            
            char_title=f"**Note:** {note}"
            markdown_to_docx_paragraph(doc, char_title)

        else:
            characterBundle=buildChar(char.name, char.description, char.personality, char.motivation, char.image_mime_type, note=note)
            markdown_to_docx_paragraph(doc, characterBundle)
            if char.image_mime_type != "":
                insert_image(doc, char.image_data, char.image_mime_type)
    return

def print_chapter_summary(doc, chapter):
    doc.add_heading("Chapter Status", level=2)
    doc.add_paragraph(f"{chapter.status}")
    if chapter.status != "in_progress":
        doc.add_heading("Summary", level=2)
        markdown_to_docx_paragraph(doc, chapter.summary)

def print_story_chars(doc, story):
    story_chars = StoryWithCharactersService.get_story_with_characters(story.story_id)
    firstTime=True
    for story_char in story_chars:
        if firstTime:
            doc.add_heading(f"Characters", level=2)
            firstTime=False
        characterBundle=buildChar(story_char.name, story_char.description, story_char.personality, story_char.motivation, story_char.image_mime_type, note=story_char.note)
        markdown_to_docx_paragraph(doc, characterBundle)
        if story_char.image_mime_type != "":
            insert_image(doc, story_char.image_data, story_char.image_mime_type)  